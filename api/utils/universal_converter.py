from fastapi import APIRouter, UploadFile, File, Form
from fastapi.responses import FileResponse
from PIL import Image, ImageDraw, ImageFont
from openpyxl import load_workbook, Workbook
from pdf2image import convert_from_path
from docx import Document
from pdf2docx import Converter
# from docx2pdf import convert
from fpdf import FPDF
from reportlab.pdfgen import canvas
from enum import Enum
from PyPDF2 import PdfReader
import os
import shutil
import json
import openpyxl
import pytesseract
import fitz
import textwrap
import PyPDF2
import pandas as pd



router = APIRouter()

class Format(str, Enum):
    docx = "docx"
    pdf = "pdf"
    xlsx = "xlsx"
    txt = "txt"
    jpeg = "jpeg"
    json = "json"
    

UPLOAD_DIR = "uploads"
CONVERTED_DIR = "converted"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(CONVERTED_DIR, exist_ok=True)

def detect_format(filename: str) -> str:
    ext = filename.split(".")[-1].lower()
    return ext

def wrap_text(text, font, max_width, draw):
    """Wrap text to fit within max_width."""
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(textwrap.wrap(paragraph, width=max_width, break_long_words=False))
    return lines

def convert_to_txt(input_path: str, output_path: str):
    ext = detect_format(input_path)
    
    if ext == "docx":
        doc = Document(input_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
            
    elif ext == "xlsx":
        df = pd.read_excel(input_path)
        df.to_csv(output_path, index=False, sep='\t')
        
    elif ext == "pdf":
        doc = fitz.open(input_path)  # Open the PDF
        text = ""
        for page in doc:
            text += page.get_text("text")  # Extract text from each page
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

    elif ext == "jpeg" or ext == "jpg":
        image = Image.open(input_path) # Open the image
        extracted_text = pytesseract.image_to_string(image)   # Extract text using Tesseract OCR
        with open(output_path, "w") as output_file: # Save text to the output file
            output_file.write(extracted_text)
            
    elif ext == "json":
        with open(input_path, 'r') as input_path:
            data = json.load(input_path)
        with open(output_path, 'w') as txt_file:
            for key, value in data.items():
                txt_file.write(f"{key}: {value}\n")


def convert_to_docx(input_path: str, output_path: str):
    ext = detect_format(input_path)
    
    if ext == "txt":
        doc = Document()
        with open(input_path, "r", encoding="utf-8") as f:
            doc.add_paragraph(f.read())
        doc.save(output_path)
        
    elif ext == "pdf":
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        
    elif ext == "xlsx":
        workbook = load_workbook(input_path)
        sheet = workbook.active
        doc = Document()    # Create a new DOCX document
        columns = list(sheet.columns)   # Add the table to DOCX
        table = doc.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells # Add headers to the first row (if the sheet has headers)
        for i, column_name in enumerate(sheet[1]):
            hdr_cells[i].text = str(column_name)
        for row in sheet.iter_rows(min_row=2, values_only=True): # Add data from XLSX to DOCX table # Skip header row
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if value is not None else ""
        doc.save(output_path)   # Save the DOCX file
        
    elif ext == "jpeg" or ext == "jpg":
        image = Image.open(input_path)
        text = pytesseract.image_to_string(image)   # Extract text from the image using Tesseract OCR
        doc = Document()    # Create a new DOCX document
        doc.add_paragraph(text) # Add the extracted text to the DOCX document
        doc.save(output_path)   # Save the DOCX file
        
    elif ext == "json":
        with open(input_path, "r", encoding="utf-8") as file:
            data = json.load(file)
        doc = Document()
        def add_content(data, indent=0):
            spacing = " " * (indent * 4)  # Indent with spaces for nested levels
            if isinstance(data, dict):
                doc.add_paragraph(f"{spacing}{{")  # Open dictionary
                for key, value in data.items():
                    doc.add_paragraph(f"{spacing}    '{key}': ")  # Keep dictionary style
                    add_content(value, indent + 1)  # Recursively format values
                doc.add_paragraph(f"{spacing}}}")  # Close dictionary
            elif isinstance(data, list):
                doc.add_paragraph(f"{spacing}[")  # Open list
                for item in data:
                    add_content(item, indent + 1)  # Process each list item
                doc.add_paragraph(f"{spacing}]")  # Close list
            else:
                doc.add_paragraph(f"{spacing}{repr(data)}")
        add_content(data)
        doc.save(output_path)
    
def convert_to_pdf(input_path: str, output_path: str):
    ext = detect_format(input_path)
    
    if ext == "txt":
        with open(input_path, "r", encoding="utf-8") as txt_file:
            text = txt_file.readlines()
        pdf = canvas.Canvas(output_path)
        pdf.setFont("Helvetica", 12)
        y = 800  # Starting position for text
        for line in text:
            if y < 50:  # Prevent text from going off the page
                pdf.showPage()
                pdf.setFont("Helvetica", 12)
                y = 800
            pdf.drawString(50, y, line.strip())
            y -= 20  # Move down for the next line
        pdf.save()
    
    elif ext == "docx":
        doc = Document(input_path)
        c = canvas.Canvas(output_path)

        y = 750  # Start writing from the top
        for para in doc.paragraphs:
            if y < 50:  # Prevent text from going off the page
                c.showPage()
                c.setFont("Helvetica", 12)
                y = 750
            c.drawString(50, y, para.text)
            y -= 20  # Move down for the next line
        c.save()
        
    elif ext == "jpeg" or ext == "jpg":
        image = Image.open(input_path)
        image.convert("RGB").save(output_path, "PDF")
        
    elif ext == "xlsx":
        workbook = openpyxl.load_workbook(input_path)
        sheet = workbook.active 
        pdf = FPDF()    # Create PDF object
        pdf.add_page()
        pdf.set_font("Courier", size=10)
        # Loop through rows and columns of the sheet and add content to PDF
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
            pdf.cell(200, 10, txt=row_text, ln=True)
        pdf.output(output_path)     # Save the PDF
        
    elif ext == "json":
        with open(input_path, "r", encoding="utf-8") as file:
            data = json.load(file)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Courier", size=12)
        pdf.multi_cell(0, 10, json.dumps(data, indent=2))  # Directly dumps JSON as text
        pdf.output(output_path)
        print(f"✅ JSON converted to PDF: {output_path}")

def convert_to_xlsx(input_path: str, output_path: str):
    ext = detect_format(input_path)
    if ext == "txt":
        df = pd.read_csv(input_path, delimiter='\t')
        df.to_excel(output_path, index=False)
        
    elif ext == "docx":
        doc = Document(input_path)
        wb = Workbook()     # Create a new Workbook for XLSX
        ws = wb.active
        ws.title = "Converted Text"
        row = 1     # Add the DOCX text to the XLSX file
        for para in doc.paragraphs:
            ws[f"A{row}"] = para.text
            row += 1
        for table in doc.tables:    # Add tables (if any) from DOCX to XLSX
            row += 1        # Start inserting data after paragraphs to avoid overlap
            for i, table_row in enumerate(table.rows):
                for j, cell in enumerate(table_row.cells):
                    ws.cell(row=row + i, column=j + 1, value=cell.text)
        wb.save(output_path)    # Save the XLSX file
        
    elif ext == "pdf":
        with open(input_path, "rb") as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            pdf_text = ""
            for page in reader.pages:
                pdf_text += page.extract_text()
        wb = Workbook()     # Create a new Workbook for XLSX
        ws = wb.active
        ws.title = "PDF Text"
        # Write the extracted text to the XLSX file
        for i, line in enumerate(pdf_text.split("\n")):
            ws[f"A{i + 1}"] = line
        # Save the XLSX file
        wb.save(output_path)
        
    elif ext == "jpg" or ext == "jpeg":
        img = Image.open(input_path)
        # Use pytesseract to extract text from the image
        extracted_text = pytesseract.image_to_string(img)
        wb = Workbook()     # Create a new Workbook for XLSX
        ws = wb.active
        ws.title = "Image Text"
        for i, line in enumerate(extracted_text.split("\n")):       # Write the extracted text to the XLSX file
            ws[f"A{i + 1}"] = line
        # Save the XLSX file
        wb.save(output_path)
        
    elif ext == "json":
        with open(input_path, "r", encoding="utf-8") as file:
            data = json.load(file)
        df = pd.DataFrame([data])
        df.to_excel(output_path, index=False)
        
def convert_to_json(input_path: str, output_path: str):
    ext = detect_format(input_path)
    
    if ext == "docx":
        doc = Document(input_path)
        data = {}
        for para in doc.paragraphs:
            data[para.text] = ""
        with open(output_path, 'w') as json_file:
            json.dump(data, json_file, indent=4)
            
    elif ext == "xlsx":
        data = pd.read_excel(input_path).to_dict(orient='records')
        with open(output_path, 'w') as json_file:
            json.dump(data, json_file, indent=4)
            
    elif ext == "pdf":
        reader = PdfReader(input_path)
        data = {}
        for page_num, page in enumerate(reader.pages):
            data[f"Page {page_num + 1}"] = page.extract_text()
        with open(output_path, 'w') as json_file:
            json.dump(data, json_file, indent=4)
            
    elif ext == "txt":
        with open(input_path, 'r') as txt_file:
            data = txt_file.readlines()
        with open(output_path, 'w') as json_file:
            json.dump(data, json_file, indent=4)
            
    elif ext == "jpeg":
        img = Image.open(input_path)
        text = pytesseract.image_to_string(img)
        data = {"extracted_text": text}
        with open(output_path, 'w') as json_file:
            json.dump(data, json_file, indent=4)

def convert_to_jpeg(input_path: str, output_path: str):
    ext = detect_format(input_path)
    
    if ext == "docx":
        doc = Document(input_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        img = Image.new('RGB', (800, 1000), color='white')
        draw = ImageDraw.Draw(img)
        font = ImageFont.load_default()
        wrapped_text = wrap_text(text, font, 150, draw)
        y_text = 40
        for line in wrapped_text:
            draw.text((20, y_text), line, fill='black', font=font)
            y_text += 20  # Adjust line spacing
        img.save(output_path)
     
    elif ext == "xlsx":
        df = pd.read_excel(input_path)
        text = df.to_string(index=False)
        img = Image.new('RGB', (1000, 800), color='white')
        draw = ImageDraw.Draw(img)
        font = ImageFont.load_default()
        wrapped_text = wrap_text(text, font, 100, draw)
        y_text = 40
        for line in wrapped_text:
            draw.text((20, y_text), line, fill='black', font=font)
            y_text += 20
        img.save(output_path)
        
    elif ext == "pdf":
        # reader = PdfReader(input_path)
        # text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        # img = Image.new('RGB', (800, 1000), color='white')
        # draw = ImageDraw.Draw(img)
        # font = ImageFont.load_default()
        # draw.text((10, 10), text[:4000], fill='black', font=font)  # Truncate long text
        # img.save(output_path)
        
        images = convert_from_path(input_path, dpi=200, fmt="jpeg")
        # Save each page as a JPEG file
        for i, image in enumerate(images):
            # output_path = os.path.join(output_folder, f"page_{i + 1}.jpeg")
            image.save(output_path, "JPEG")
        
        # images = convert_from_path(input_path, dpi=200, fmt="jpeg")
        
        # # Get the directory and filename parts of output_path
        # output_dir = os.path.dirname(output_path)
        # filename_base = os.path.splitext(os.path.basename(output_path))[0]
        # # Create output directory if it doesn't exist
        # if output_dir and not os.path.exists(output_dir):
        #     os.makedirs(output_dir)
        # # List to store all output paths
        # output_paths = []
        # # Save each image
        # for i, image in enumerate(images):
        #     # For single page PDF, use the exact output_path
        #     if len(images) == 1:
        #         current_output_path = output_path
        #     else:
        #         # For multi-page PDF, add page number to filename
        #         current_output_path = os.path.join(
        #             output_dir,
        #             f"{filename_base}_page_{i+1}.jpg"
        #         )
        #     # Save the image
        #     image.save(current_output_path, "JPEG")
        #     output_paths.append(current_output_path)
        
        # return output_paths
        
    elif ext == "txt":
        with open(input_path, 'r') as file:
            text = file.read()
        img = Image.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(img)
        font = ImageFont.load_default()
        draw.text((10, 10), text[:4000], fill='black', font=font)  # Prevent overflow
        img.save(output_path)
        
    elif ext == "json":
        with open(input_path, 'r') as file:
            data = json.load(file)
        text = json.dumps(data, indent=4)
        img = Image.new('RGB', (1000, 800), color='white')
        draw = ImageDraw.Draw(img)
        font = ImageFont.load_default()
        draw.text((10, 10), text[:4000], fill='black', font=font)  # Prevent overflow
        img.save(output_path)
        
    

@router.post("/convert")
async def convert_file(file: UploadFile = File(...), output_format: Format = Form(...)):
    input_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    ext = detect_format(file.filename)
    output_filename = f"{file.filename}.{output_format}"
    output_path = os.path.join(CONVERTED_DIR, output_filename)
    
    if output_format == "txt":
        convert_to_txt(input_path, output_path)
    elif output_format == "docx":
        convert_to_docx(input_path, output_path)
    elif output_format == "pdf":
        convert_to_pdf(input_path, output_path)
    elif output_format == "xlsx":
        convert_to_xlsx(input_path, output_path)
    elif output_format == "json":
        convert_to_json(input_path, output_path)
    elif output_format == "jpeg":
        convert_to_jpeg(input_path, output_path)
    else:
        return {"error": "Unsupported format"}
    
    return FileResponse(output_path, filename=output_filename)

